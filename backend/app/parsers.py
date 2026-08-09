# -*- coding: utf-8 -*-
"""文档解析：txt / md / docx / pdf → 纯文本（M3）。

解析失败抛 ValueError（上层捕获后置 failed + error_message）。
"""
import logging
from pathlib import Path

import chardet

logger = logging.getLogger(__name__)

SUPPORTED_EXT = {"txt", "md", "docx", "pdf"}


def parse_file(path: Path) -> str:
    """按扩展名解析文档为纯文本。"""
    ext = path.suffix.lower().lstrip(".")
    if ext not in SUPPORTED_EXT:
        raise ValueError(f"不支持的文件格式: {ext}")
    if ext == "txt":
        return _parse_txt(path)
    if ext == "md":
        return _parse_txt(path)
    if ext == "docx":
        return _parse_docx(path)
    return _parse_pdf(path)


def _parse_txt(path: Path) -> str:
    raw = path.read_bytes()
    encoding = chardet.detect(raw)["encoding"] or "utf-8"
    try:
        return raw.decode(encoding, errors="replace")
    except LookupError:  # 编码名无效时回退
        return raw.decode("utf-8", errors="replace")


def _parse_docx(path: Path) -> str:
    import docx  # python-docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:  # 表格内容也提取（换行拼接）
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(path: Path) -> str:
    import fitz  # PyMuPDF

    pages = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                pages.append(text)
    if not pages:
        raise ValueError("PDF 无文本内容（可能为扫描件）")
    return "\n".join(pages)
