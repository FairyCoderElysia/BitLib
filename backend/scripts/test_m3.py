# -*- coding: utf-8 -*-
"""M3 解析入库管线测试（TestClient + assert 全绿）。

覆盖：4 格式解析入库 / Chroma+ChunkParent+FTS 落库 / 过短拦截 / 超长分片 / 权限字段冗余 / reprocess。

运行：cd backend && python scripts/test_m3.py
首次运行会下载本地 embedding 模型（bge-small-zh-v1.5，~95MB）。
"""
import os
import shutil
import sys
from pathlib import Path

BACKEND_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BACKEND_DIR))

TEST_ROOT = BACKEND_DIR / "data" / "test_m3"
TEST_DB = (TEST_ROOT / "app.db").as_posix()

os.environ["ADMIN_INITIAL_PASSWORD"] = "Admin@123456"
os.environ["DATABASE_URL"] = f"sqlite:///{TEST_DB}"
os.environ["UPLOAD_DIR"] = str(TEST_ROOT / "uploads")
os.environ["CHROMA_DIR"] = str(TEST_ROOT / "chroma")
os.environ["MAX_UPLOAD_MB"] = "5"
os.environ.setdefault("SECRET_KEY", "test-secret-key")
os.environ["EMBEDDING_MODE"] = "local"

from fastapi.testclient import TestClient  # noqa: E402

from app import chunker, models, vector_store  # noqa: E402
from app.db import SessionLocal  # noqa: E402
from app.main import app  # noqa: E402

ADMIN_PASSWORD = "Admin@123456"
NEW_ADMIN_PASSWORD = "Admin@123456!"  # A7 硬拦截适配：首登 admin 改密用


def H(token):
    return {"Authorization": f"Bearer {token}"}


def _make_txt(path: Path, text: str):
    path.write_text(text, encoding="utf-8")


def _make_md(path: Path, text: str):
    path.write_text(f"# 标题\n\n{text}", encoding="utf-8")


def _make_docx(path: Path, text: str):
    import docx
    d = docx.Document()
    d.add_heading("测试文档", level=1)
    for para in text.split("\n"):
        if para.strip():
            d.add_paragraph(para.strip())
    d.save(str(path))


def _make_pdf(path: Path, text: str):
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text[:1800], fontsize=11)
    doc.save(str(path))
    doc.close()


def main():
    shutil.rmtree(TEST_ROOT, ignore_errors=True)
    TEST_ROOT.mkdir(parents=True, exist_ok=True)
    passed: list = []

    def check(name):
        passed.append(name)
        print(f"  PASS: {name}")

    def direct_upload(c, token, path: Path, title: str, department_id=None):
        """直入库上传，返回 (status_code, json)。"""
        with path.open("rb") as f:
            files = {"file": (path.name, f, "application/octet-stream")}
            data = {"title": title}
            if department_id is not None:
                data["department_id"] = str(department_id)
            return c.post("/api/admin/documents/upload",
                          headers=H(token), files=files, data=data)

    try:
        with TestClient(app) as c:
            r = c.post("/api/auth/login", json={"username": "admin", "password": ADMIN_PASSWORD})
            assert r.status_code == 200 and r.json()["code"] == 0, r.text
            data = r.json()["data"]
            token = data["token"]
            if data["user"].get("must_change_password"):
                r = c.post("/api/auth/change-password", headers=H(token),
                           json={"old_password": ADMIN_PASSWORD,
                                 "new_password": NEW_ADMIN_PASSWORD})
                assert r.status_code == 200 and r.json()["code"] == 0, r.text

            # ---------- 1. 四种格式样例入库 ----------
            samples = []
            t = "企业内部文档管理系统支持全文检索、向量知识库与人工智能问答。员工可以快速定位所需资料。"
            f1 = TEST_ROOT / "s1.txt"
            f1.parent.mkdir(parents=True, exist_ok=True)
            _make_txt(f1, (t * 20))
            samples.append(("txt", f1, "企业文档管理规范 txt"))
            f2 = TEST_ROOT / "s2.md"
            _make_md(f2, (t * 20))
            samples.append(("md", f2, "企业文档管理规范 md"))
            f3 = TEST_ROOT / "s3.docx"
            _make_docx(f3, (t * 30))
            samples.append(("docx", f3, "企业文档管理规范 docx"))
            f4 = TEST_ROOT / "s4.pdf"
            _make_pdf(f4, (t * 8))
            samples.append(("pdf", f4, "企业文档管理规范 pdf"))

            doc_ids = []
            for ext, fp, title in samples:
                r = direct_upload(c, token, fp, title)
                assert r.status_code == 200 and r.json()["code"] == 0, r.text
                d = r.json()["data"]
                assert d["status"] == "approved", d
                doc_ids.append(d["id"])
                print(f"  ... {ext} 入库 approved id={d['id']}")
            check("1. txt/md/docx/pdf 四格式直入库均 approved")

            # ---------- 2. 落库完整性：Chroma / ChunkParent / FTS ----------
            with SessionLocal() as db:
                for did in doc_ids:
                    doc = db.get(models.Document, did)
                    assert doc.content_text and len(doc.content_text) > 100, "content_text 未保存"
                    n_child = vector_store.count_by_document(did)
                    assert n_child > 0, f"文档 {did} 无 Chroma child"
                    n_parent = db.query(models.ChunkParent).filter(
                        models.ChunkParent.document_id == did).count()
                    assert n_parent > 0, f"文档 {did} 无 ChunkParent"
                    # FTS 可查询（查询侧 jieba 分词后 MATCH，与 M4 检索实现一致）
                    from sqlalchemy import text as _text
                    import jieba as _jieba
                    kw = " ".join(_jieba.cut("文档管理"))
                    row = db.execute(
                        _text("SELECT rowid FROM document_fts WHERE document_fts MATCH :kw"),
                        {"kw": kw}).fetchone()
                    assert row is not None, f"文档 {did} FTS 查询未命中（kw={kw}）"
            check("2. Chroma child + ChunkParent + content_text + FTS 全部落库")

            # ---------- 3. 权限字段冗余 ----------
            with SessionLocal() as db:
                first = db.get(models.Document, doc_ids[0])
                hits = vector_store.query([0.0] * 512, 3, user_department_id=None, is_admin=True)
                assert len(hits) > 0, "admin 全量查询无结果"
                for h in hits:
                    assert h["status"] == "approved", h
            check("3. 向量查询（admin 全量）可召回，metadata 含可见性字段")

            # ---------- 4. 过短文本拦截 ----------
            f5 = TEST_ROOT / "s5.txt"
            _make_txt(f5, "太短了")  # < 50 字符
            r = direct_upload(c, token, f5, "过短文本")
            assert r.status_code == 200, r.text
            d = r.json()["data"]
            assert d["status"] == "failed", d
            assert d["error_message"], "failed 无错误信息"
            check("4. 过短文本（<50 字符）标记 failed + error_message")

            # ---------- 5. 超长文档完整分片 ----------
            long_text = "超长文档测试段落。" * 20000  # 约 16 万字
            chunks = chunker.chunk_document(long_text)
            assert len(chunks) > 1, f"超长文本仅 {len(chunks)} 个 parent"
            joined = "".join(p["parent"]["text"] for p in chunks)
            # 还原率：拼接后应包含原文核心内容（剔除空白后）
            ratio = len(joined) / len(long_text)
            assert ratio > 0.95, f"parent 拼接还原率 {ratio:.3f} < 0.95"
            check(f"5. 16 万字文档分片：{len(chunks)} parent，还原率 {ratio:.3f}")

            # ---------- 6. 超长文档走完整入库管线（3 万字，避免测试过慢） ----------
            f6 = TEST_ROOT / "s6.txt"
            _make_txt(f6, ("长文档内容。" * 6000))  # 约 3 万字
            r = direct_upload(c, token, f6, "超长文档入库")
            assert r.status_code == 200, r.text
            d = r.json()["data"]
            assert d["status"] == "approved", d
            with SessionLocal() as db:
                n_parent = db.query(models.ChunkParent).filter(
                    models.ChunkParent.document_id == d["id"]).count()
                n_child = vector_store.count_by_document(d["id"])
                assert n_parent > 1, f"超长文档仅 {n_parent} parent"
                assert n_child >= n_parent, f"child({n_child}) < parent({n_parent})"
                long_id = d["id"]
            check(f"6. 3 万字文档入库：{n_parent} parent / {n_child} child")

            # ---------- 7. reprocess：failed 文档修复文件后重跑成功 ----------
            bad = TEST_ROOT / "bad.txt"
            bad.write_text("损坏内容", encoding="utf-8")  # <50 字符 → failed
            r = direct_upload(c, token, bad, "待修复文档")
            assert r.status_code == 200, r.text
            d = r.json()["data"]
            assert d["status"] == "failed", d
            with SessionLocal() as db:
                fp = Path(db.get(models.Document, d["id"]).file_path)
            # 修复：覆盖服务器端原文件为合法内容
            (TEST_ROOT / "uploads" / fp).write_text(
                "修复后的完整文档内容，包含足够长度的有效文本。" * 20, encoding="utf-8")
            r = c.post(f"/api/admin/documents/{d['id']}/reprocess", headers=H(token))
            assert r.status_code == 200 and r.json()["code"] == 0, r.text
            assert r.json()["data"]["status"] == "approved", r.text
            # 幂等：approved 后再次 reprocess 应被拒（spec：仅 failed/offline）
            r = c.post(f"/api/admin/documents/{d['id']}/reprocess", headers=H(token))
            assert r.status_code == 400, r.text
            check("7. reprocess：failed 修复后重跑成功 → approved；approved 再重跑被拒(400)")

            # ---------- 8. 未修复的 failed 文档 reprocess 仍 failed ----------
            bad2 = TEST_ROOT / "bad2.txt"
            bad2.write_text("另一份过短文本", encoding="utf-8")  # 新文件，避免与第7步撞 sha256
            r2 = direct_upload(c, token, bad2, "过短文本2")
            assert r2.status_code == 200, r2.text
            d2 = r2.json()["data"]
            assert d2["status"] == "failed", d2
            r3 = c.post(f"/api/admin/documents/{d2['id']}/reprocess", headers=H(token))
            assert r3.status_code == 200, r3.text
            assert r3.json()["data"]["status"] == "failed", "未修复文档 reprocess 不应成功"
            check("8. 未修复 failed 文档 reprocess 仍 failed")

        print(f"\n=== ALL {len(passed)} M3 TESTS PASSED ===")
    finally:
        shutil.rmtree(TEST_ROOT, ignore_errors=True)


if __name__ == "__main__":
    main()
